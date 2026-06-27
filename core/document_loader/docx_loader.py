from docx import Document
from core.document_loader.base import DocumentLoader
from core.document_loader.cleaner import clean_text
from utils.logger import logger


class DocxLoader(DocumentLoader):
    def load(self) -> str:
        doc = Document(self.file_path)
        text_parts = []
        logger.info(f"解析DOCX文档: {self.file_path.name}")

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = clean_text("\n".join(text_parts))
        logger.info(
            f"DOCX解析完成: {self.file_path.name}, "
            f"提取字符数: {len(full_text)}"
        )
        return full_text

    def get_metadata(self) -> dict:
        doc = Document(self.file_path)
        return {
            "source_doc": self.file_path.name,
            "file_type": ".docx",
            "file_size": self.file_path.stat().st_size,
            "paragraph_count": len(doc.paragraphs),
        }
